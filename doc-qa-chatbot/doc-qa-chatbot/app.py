import logging
import os
import tempfile
import warnings

# Suppress Hugging Face, Transformers & PyTorch Warning Logs
os.environ["TRANSFORMERS_NO_ADVISORY_WARNINGS"] = "1"
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"
warnings.filterwarnings("ignore")
logging.getLogger("transformers").setLevel(logging.ERROR)
logging.getLogger("huggingface_hub").setLevel(logging.ERROR)

import docx
import ollama
import pdfplumber
import streamlit as st
import streamlit.components.v1 as components
from pdf2image import convert_from_path
import pytesseract

from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

# ------------------------------------------------------------------------------
# 1. UI Setup & 3D WebGL Canvas
# ------------------------------------------------------------------------------
st.set_page_config(
    page_title="Multi-File Document Hub", page_icon="📂", layout="wide"
)

st.markdown(
    """
    <style>
    .stApp { background-color: #0b0f19; color: #f1f5f9; }
    .stButton>button {
        background: linear-gradient(135deg, #6366f1 0%, #a855f7 100%);
        color: white; border: none; border-radius: 8px; font-weight: 600;
    }
    </style>
""",
    unsafe_allow_html=True,
)

three_js_header = """
<!DOCTYPE html>
<html>
<head>
    <style> body { margin: 0; overflow: hidden; background: #0b0f19; } #canvas-container { width: 100vw; height: 180px; } </style>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/three.js/r128/three.min.js"></script>
</head>
<body>
    <div id="canvas-container"></div>
    <script>
        const container = document.getElementById('canvas-container');
        const scene = new THREE.Scene();
        const camera = new THREE.PerspectiveCamera(60, container.clientWidth / container.clientHeight, 0.1, 1000);
        camera.position.z = 2.8;
        const renderer = new THREE.WebGLRenderer({ antialias: true, alpha: true });
        renderer.setSize(container.clientWidth, container.clientHeight);
        container.appendChild(renderer.domElement);
        
        const count = 1500;
        const geometry = new THREE.BufferGeometry();
        const pos = new Float32Array(count * 3);
        for(let i = 0; i < count * 3; i += 3) {
            pos[i] = (Math.random() - 0.5) * 5;
            pos[i+1] = (Math.random() - 0.5) * 5;
            pos[i+2] = (Math.random() - 0.5) * 5;
        }
        geometry.setAttribute('position', new THREE.BufferAttribute(pos, 3));
        const mat = new THREE.PointsMaterial({ size: 0.02, color: 0x818cf8, transparent: true, opacity: 0.7 });
        const particles = new THREE.Points(geometry, mat);
        scene.add(particles);

        function animate() {
            requestAnimationFrame(animate);
            particles.rotation.y += 0.002;
            renderer.render(scene, camera);
        }
        animate();
    </script>
</body>
</html>
"""
components.html(three_js_header, height=180)

st.title("📂 Multi-File & Folder Intelligence Hub")

# ------------------------------------------------------------------------------
# 2. File Parsing & Processing Functions
# ------------------------------------------------------------------------------
tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
if os.path.exists(tesseract_path):
    pytesseract.pytesseract.tesseract_cmd = tesseract_path


def parse_single_file(file_obj):
    filename = file_obj.name
    ext = os.path.splitext(filename)[1].lower()
    docs = []

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(file_obj.getvalue())
        tmp_path = tmp.name

    try:
        if ext == ".pdf":
            try:
                loader = PyPDFLoader(tmp_path)
                docs = loader.load()
            except Exception:
                docs = []

            if not docs or not any(d.page_content.strip() for d in docs):
                with pdfplumber.open(tmp_path) as pdf:
                    for i, page in enumerate(pdf.pages):
                        t = page.extract_text()
                        if t and t.strip():
                            docs.append(
                                Document(
                                    page_content=t,
                                    metadata={
                                        "source": filename,
                                        "page": i + 1,
                                    },
                                )
                            )

        elif ext in [".txt", ".md"]:
            text = file_obj.getvalue().decode("utf-8", errors="ignore")
            if text.strip():
                docs.append(
                    Document(page_content=text, metadata={"source": filename})
                )

        elif ext == ".docx":
            doc = docx.Document(tmp_path)
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text])
            if full_text.strip():
                docs.append(
                    Document(
                        page_content=full_text, metadata={"source": filename}
                    )
                )

    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

    for d in docs:
        d.metadata["source"] = filename
    return docs


@st.cache_resource(show_spinner="Processing and indexing batch files...")
def build_vectorstore(uploaded_files):
    all_docs = []
    for file_obj in uploaded_files:
        parsed_docs = parse_single_file(file_obj)
        all_docs.extend(parsed_docs)

    if not all_docs:
        raise ValueError("No valid text content found in uploaded files.")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, chunk_overlap=200
    )
    splits = splitter.split_documents(all_docs)

    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    return Chroma.from_documents(documents=splits, embedding=embeddings)


# ------------------------------------------------------------------------------
# 3. Sidebar Multi-File Upload & Chat UI
# ------------------------------------------------------------------------------
with st.sidebar:
    st.header("⚙️ Workspace Controls")
    uploaded_files = st.file_uploader(
        "Deploy Files or Folders",
        type=["pdf", "txt", "docx", "md"],
        accept_multiple_files=True,
        key="file_batch",
    )

    if uploaded_files:
        st.write(f"📁 **Staged Items:** {len(uploaded_files)} file(s)")

    if st.button("🗑️ Clear Chat Context", use_container_width=True):
        st.session_state.messages = []
        st.rerun()

if uploaded_files:
    try:
        vectorstore = build_vectorstore(uploaded_files)
        st.success(f"✅ Indexed {len(uploaded_files)} file(s) into ChromaDB.")

        if "messages" not in st.session_state:
            st.session_state.messages = []

        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])

        user_query = st.chat_input("Ask a question across all deployed files...")

        if user_query:
            st.session_state.messages.append(
                {"role": "user", "content": user_query}
            )
            with st.chat_message("user"):
                st.write(user_query)

            results = vectorstore.similarity_search(user_query, k=4)
            context = "\n\n".join([doc.page_content for doc in results])

            with st.chat_message("assistant"):
                with st.spinner("Analyzing uploaded knowledge base..."):
                    try:
                        response = ollama.chat(
                            model="llama3.2",
                            messages=[
                                {
                                    "role": "system",
                                    "content": "Answer strictly based on the provided document context.",
                                },
                                {
                                    "role": "user",
                                    "content": f"Context:\n{context}\n\nQuestion: {user_query}",
                                },
                            ],
                        )

                        answer = response["message"]["content"]
                        st.write(answer)

                        with st.expander("🔍 View Context Sources & Files"):
                            for idx, doc in enumerate(results, 1):
                                src = doc.metadata.get("source", "Unknown")
                                page = doc.metadata.get("page", "N/A")
                                st.markdown(
                                    f"**Chunk {idx}** | Source: `{src}` (Page: {page})"
                                )
                                st.caption(doc.page_content[:300] + "...")

                        st.session_state.messages.append(
                            {"role": "assistant", "content": answer}
                        )
                    except Exception as err:
                        st.error(f"Ollama execution error: {str(err)}")

    except Exception as e:
        st.error(f"Error processing files: {str(e)}")
else:
    st.info("👈 Drag and drop files or entire folders into the sidebar upload zone.")